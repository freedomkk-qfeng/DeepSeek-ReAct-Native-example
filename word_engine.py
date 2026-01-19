import os
from docx import Document


class WordEngine:
    """
    Word文档操作引擎，专门设计用于支撑AI代理。
    所有表格操作工具都是通用的，不包含任何硬编码的业务逻辑。
    """
    def __init__(self, file_path=None):
        self.file_path = file_path
        if file_path and os.path.exists(file_path):
            self.doc = Document(file_path)
            print(f"已加载文档: {file_path}")
        else:
            self.doc = Document()
            print("已创建新文档")

    # ==================== 通用表格操作工具 ====================
    
    def _get_grid_span(self, cell):
        """获取单元格的列跨度"""
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is not None:
            gridSpan = tcPr.gridSpan
            if gridSpan is not None:
                return gridSpan.val
        return 1
    
    def analyze_table(self, table_index=0):
        """
        深度分析表格结构，返回完整的表格视图，供AI理解和决策。
        
        不做任何假设，让AI自己判断：
        - 哪些是标签，哪些是值
        - 表格是键值对布局还是列表布局
        - 哪些单元格需要填写
        
        返回包含：
        - 表格基本信息
        - 每行的唯一单元格（处理合并单元格，使用gridSpan）
        - 自动检测的标签-值对
        - 空单元格列表
        """
        try:
            if table_index >= len(self.doc.tables):
                return {"error": f"表格索引 {table_index} 超出范围，文档共有 {len(self.doc.tables)} 个表格"}
            
            table = self.doc.tables[table_index]
            result = {
                "table_index": table_index,
                "total_rows": len(table.rows),
                "total_cols": len(table.columns),
                "rows": [],
                "fillable_cells": [],
                "label_value_pairs": []
            }
            
            for r, row in enumerate(table.rows):
                # 使用 gridSpan 正确处理合并单元格
                unique_cells = []
                c = 0
                while c < len(row.cells):
                    cell = row.cells[c]
                    span = self._get_grid_span(cell)
                    cell_text = cell.text.strip().replace('\n', ' ')
                    unique_cells.append({
                        "col": c,
                        "span": span,
                        "text": cell_text,
                        "is_empty": len(cell_text) == 0
                    })
                    c += span
                
                row_info = {
                    "row": r,
                    "unique_cells": unique_cells,
                    "cell_count": len(unique_cells)
                }
                result["rows"].append(row_info)
                
                # 自动识别键值对模式：
                # 如果一个非空短文本后面跟着空单元格，可能是标签-值对
                i = 0
                while i < len(unique_cells) - 1:
                    current = unique_cells[i]
                    next_cell = unique_cells[i + 1]
                    
                    # 当前单元格像标签（非空、较短）
                    if not current["is_empty"] and len(current["text"]) <= 20:
                        # 计算值单元格的实际列位置
                        value_col = current["col"] + current["span"]
                        
                        if next_cell["is_empty"]:
                            # 后面是空单元格，这是需要填写的键值对
                            result["label_value_pairs"].append({
                                "label": current["text"],
                                "label_position": {"row": r, "col": current["col"]},
                                "value_position": {"row": r, "col": value_col},
                                "current_value": "",
                                "needs_fill": True
                            })
                            result["fillable_cells"].append({
                                "row": r,
                                "col": value_col,
                                "associated_label": current["text"]
                            })
                            i += 2
                            continue
                        elif len(next_cell["text"]) > 0:
                            # 后面有值，这是已填写的键值对
                            result["label_value_pairs"].append({
                                "label": current["text"],
                                "label_position": {"row": r, "col": current["col"]},
                                "value_position": {"row": r, "col": value_col},
                                "current_value": next_cell["text"],
                                "needs_fill": False
                            })
                            i += 2
                            continue
                    
                    # 独立的空单元格
                    if current["is_empty"]:
                        result["fillable_cells"].append({
                            "row": r,
                            "col": current["col"],
                            "associated_label": None
                        })
                    i += 1
            
            return result
        except Exception as e:
            return {"error": str(e)}
    
    def get_table_as_text(self, table_index=0):
        """
        将表格转换为易读的文本格式，便于AI阅读和理解表格内容。
        
        格式：每行显示 [列号]内容，用 | 分隔。
        空单元格显示为 (空)。
        处理合并单元格，只显示唯一内容。
        """
        try:
            if table_index >= len(self.doc.tables):
                return f"错误: 表格索引 {table_index} 超出范围"
            
            table = self.doc.tables[table_index]
            lines = [f"表格 {table_index}: {len(table.rows)} 行 x {len(table.columns)} 列"]
            lines.append("-" * 50)
            
            for r, row in enumerate(table.rows):
                unique_texts = []
                seen_tc = set()
                for c, cell in enumerate(row.cells):
                    tc_id = id(cell._tc)
                    if tc_id not in seen_tc:
                        text = cell.text.strip().replace('\n', ' ')
                        if not text:
                            text = "(空)"
                        unique_texts.append(f"[{c}]{text}")
                        seen_tc.add(tc_id)
                lines.append(f"Row {r}: " + " | ".join(unique_texts))
            
            return "\n".join(lines)
        except Exception as e:
            return f"错误: {str(e)}"
    
    def list_all_tables(self):
        """
        列出文档中所有表格的概要信息。
        返回每个表格的索引、行列数和首行预览。
        """
        if not self.doc.tables:
            return "文档中没有表格"
        
        result = []
        for i, table in enumerate(self.doc.tables):
            preview = ""
            if table.rows:
                first_row_texts = []
                seen_tc = set()
                for cell in table.rows[0].cells:
                    tc_id = id(cell._tc)
                    if tc_id not in seen_tc:
                        text = cell.text.strip().replace('\n', ' ')[:15]
                        if text:
                            first_row_texts.append(text)
                        seen_tc.add(tc_id)
                preview = " | ".join(first_row_texts[:4])
                if len(first_row_texts) > 4:
                    preview += " ..."
            
            result.append({
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "preview": preview
            })
        return result
    
    def fill_cell(self, table_index, row, col, value):
        """
        填写指定位置的单元格。
        
        参数:
        - table_index: 表格索引
        - row: 行号（从0开始）
        - col: 列号（从0开始）
        - value: 要填入的值
        """
        try:
            if table_index >= len(self.doc.tables):
                return f"错误: 表格索引 {table_index} 超出范围"
            
            table = self.doc.tables[table_index]
            if row >= len(table.rows):
                return f"错误: 行号 {row} 超出范围，表格共 {len(table.rows)} 行"
            if col >= len(table.rows[row].cells):
                return f"错误: 列号 {col} 超出范围"
            
            table.rows[row].cells[col].text = str(value)
            return f"成功: 表格{table_index} 第{row}行第{col}列 已填入 '{value}'"
        except Exception as e:
            return f"填写失败: {str(e)}"
    
    def fill_by_label(self, table_index, label, value, search_all_cols=True, fill_all=False):
        """
        根据标签文本查找并填写其关联的值单元格。
        
        这是最智能的填写方法：
        1. 在表格中搜索包含指定标签文本的单元格
        2. 找到标签后，计算其列跨度，定位右侧的值单元格
        3. 将值填入该单元格
        
        参数:
        - table_index: 表格索引
        - label: 标签文本（如"姓名"、"日期"、"金额"等任意标签）
        - value: 要填入的值
        - search_all_cols: 是否搜索所有列（默认True）
        - fill_all: 是否填写所有匹配项（默认False，只填第一个）
        
        注意: 此方法完全通用，不假设任何特定的标签内容。
        """
        try:
            if table_index >= len(self.doc.tables):
                return f"错误: 表格索引 {table_index} 超出范围"
            
            table = self.doc.tables[table_index]
            label_clean = label.replace(' ', '').replace('\u3000', '')
            
            # 搜索包含标签的单元格
            # 使用 (row, col) 位置去重，而不是 tc_id
            found_positions = []
            processed_positions = set()
            
            for r, row in enumerate(table.rows):
                cols_to_search = range(len(row.cells)) if search_all_cols else [0]
                c = 0
                while c < len(row.cells):
                    if c not in cols_to_search:
                        c += 1
                        continue
                    
                    cell = row.cells[c]
                    cell_text = cell.text.strip().replace(' ', '').replace('\u3000', '')
                    
                    # 跳过空单元格
                    if not cell_text:
                        c += 1
                        continue
                    
                    # 检查是否匹配标签
                    if label_clean == cell_text or label_clean in cell_text or cell_text in label_clean:
                        pos_key = (r, c)
                        if pos_key not in processed_positions:
                            span = self._get_grid_span(cell)
                            found_positions.append((r, c, span, cell.text.strip()))
                            processed_positions.add(pos_key)
                            # 跳过合并的列
                            c += span
                            continue
                    
                    c += 1
            
            if not found_positions:
                return f"未找到包含 '{label}' 的单元格"
            
            # 填写值单元格
            positions_to_fill = found_positions if fill_all else [found_positions[0]]
            filled_count = 0
            filled_positions = set()
            
            for r, label_col, label_span, original_label in positions_to_fill:
                row_cells = table.rows[r].cells
                # 值单元格的起始位置 = 标签起始列 + 标签跨度
                value_col = label_col + label_span
                
                if value_col < len(row_cells):
                    pos_key = (r, value_col)
                    if pos_key not in filled_positions:
                        row_cells[value_col].text = str(value)
                        filled_positions.add(pos_key)
                        filled_count += 1
            
            if filled_count > 0:
                return f"成功: 根据标签 '{label}' 填入 '{value}'"
            else:
                return f"找到标签 '{label}' 但未能定位其值单元格"
        except Exception as e:
            return f"填写失败: {str(e)}"
    
    def fill_multiple_by_labels(self, table_index, label_value_map, fill_all=False):
        """
        批量根据标签填写多个值。
        
        参数:
        - table_index: 表格索引
        - label_value_map: 字典，键为标签文本，值为要填入的内容
        - fill_all: 是否填写所有匹配项（默认False，只填第一个）
        
        示例:
        fill_multiple_by_labels(0, {
            "姓名": "张三",
            "日期": "2024-01-01",
            "金额": "1000元"
        })
        """
        results = []
        for label, value in label_value_map.items():
            result = self.fill_by_label(table_index, label, value, fill_all=fill_all)
            results.append(result)
        return "\n".join(results)
    
    def find_and_fill_empty_cells_in_row(self, table_index, row_index, values, start_col=0):
        """
        在指定行中从左到右依次填写空单元格。
        
        适用于列表型表格，如明细行、项目列表等。
        
        参数:
        - table_index: 表格索引
        - row_index: 行号
        - values: 值列表，按顺序填入空单元格
        - start_col: 开始搜索的列号（默认0）
        """
        try:
            if table_index >= len(self.doc.tables):
                return f"错误: 表格索引 {table_index} 超出范围"
            
            table = self.doc.tables[table_index]
            if row_index >= len(table.rows):
                return f"错误: 行号 {row_index} 超出范围"
            
            row = table.rows[row_index]
            
            # 找出所有空的唯一单元格
            empty_cells = []
            seen_tc = set()
            for c in range(start_col, len(row.cells)):
                cell = row.cells[c]
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                if not cell.text.strip():
                    empty_cells.append((c, cell))
            
            # 依次填入值
            filled = []
            for i, value in enumerate(values):
                if i >= len(empty_cells):
                    break
                col, cell = empty_cells[i]
                cell.text = str(value)
                filled.append(f"列{col}={value}")
            
            if filled:
                return f"成功: 第{row_index}行填入 {', '.join(filled)}"
            else:
                return f"第{row_index}行没有空单元格可填写"
        except Exception as e:
            return f"填写失败: {str(e)}"
    
    def find_empty_row(self, table_index, check_col=0, start_row=1):
        """
        查找表格中第一个空行（指定列为空的行）。
        
        适用于需要追加数据的列表型表格。
        
        参数:
        - table_index: 表格索引
        - check_col: 用于判断行是否为空的列号（默认0）
        - start_row: 开始搜索的行号（默认1，跳过表头）
        
        返回: 空行的行号，如果没有则返回-1
        """
        try:
            if table_index >= len(self.doc.tables):
                return -1
            
            table = self.doc.tables[table_index]
            for r in range(start_row, len(table.rows)):
                row = table.rows[r]
                if check_col < len(row.cells):
                    if not row.cells[check_col].text.strip():
                        return r
            return -1
        except Exception:
            return -1
